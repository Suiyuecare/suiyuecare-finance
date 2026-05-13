from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.worksheet.datavalidation import DataValidation


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "assets" / "templates"
OUT.mkdir(parents=True, exist_ok=True)

XLSX_PATH = OUT / "零用金明細表.xlsx"
DOCX_PATH = OUT / "零用金明細表.docx"

ORANGE = "D97706"
LIGHT_ORANGE = "FFF4E6"
PALE = "FFF8ED"
TEXT = "1F2937"
BORDER = "E7D3B5"


def set_cell_fill(cell, fill):
    cell.fill = PatternFill("solid", fgColor=fill)


def add_border(cell):
    side = Side(style="thin", color=BORDER)
    cell.border = Border(left=side, right=side, top=side, bottom=side)


def build_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "零用金明細表"
    ws.sheet_view.showGridLines = False

    widths = {
        "A": 8,
        "B": 13,
        "C": 14,
        "D": 18,
        "E": 16,
        "F": 26,
        "G": 18,
        "H": 13,
        "I": 13,
        "J": 13,
        "K": 14,
        "L": 14,
        "M": 24,
    }
    for col, width in widths.items():
        ws.column_dimensions[col].width = width

    ws.merge_cells("A1:M1")
    ws["A1"] = "零用金明細表"
    ws["A1"].font = Font(name="Arial", bold=True, size=18, color=TEXT)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 30

    meta = [
        ("A3", "申請日期", "B3", ""),
        ("D3", "公司 / 法人", "E3", ""),
        ("G3", "組別 / 部門", "H3", ""),
        ("J3", "保管人", "K3", ""),
        ("A4", "申請類型", "B4", ""),
        ("D4", "上期餘額", "E4", ""),
        ("G4", "核定額度", "H4", ""),
        ("J4", "本次申請補足金額", "K4", "=MAX(H4-E4,0)"),
    ]
    for label_cell, label, value_cell, value in meta:
        ws[label_cell] = label
        ws[label_cell].font = Font(name="Arial", bold=True, color="7C3F00")
        set_cell_fill(ws[label_cell], LIGHT_ORANGE)
        ws[value_cell] = value
        ws[value_cell].font = Font(name="Arial", color=TEXT)
        set_cell_fill(ws[value_cell], "FFFFFF")
        add_border(ws[label_cell])
        add_border(ws[value_cell])
        ws[label_cell].alignment = Alignment(horizontal="center", vertical="center")
        ws[value_cell].alignment = Alignment(horizontal="left", vertical="center")

    ws["A6"] = "填寫說明"
    ws["A6"].font = Font(name="Arial", bold=True, color="7C3F00")
    ws["B6"] = "請逐筆填寫支出明細，含稅金額為正式認列金額；憑證照片或電子檔請一併上傳系統送簽核。"
    ws.merge_cells("B6:M6")
    ws["B6"].alignment = Alignment(wrap_text=True, vertical="center")
    set_cell_fill(ws["A6"], LIGHT_ORANGE)
    set_cell_fill(ws["B6"], PALE)
    for row in ws["A6:M6"]:
        for cell in row:
            add_border(cell)

    headers = [
        "序號",
        "日期",
        "憑證類型",
        "憑證號碼 / 發票號碼",
        "費用類別",
        "摘要 / 用途",
        "付款對象",
        "未稅",
        "營業稅",
        "含稅金額",
        "支付方式",
        "是否已附憑證",
        "備註",
    ]
    header_row = 8
    for idx, header in enumerate(headers, 1):
        cell = ws.cell(header_row, idx, header)
        cell.font = Font(name="Arial", bold=True, color="FFFFFF")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        set_cell_fill(cell, ORANGE)
        add_border(cell)
    ws.row_dimensions[header_row].height = 28

    for row in range(9, 29):
        ws.cell(row, 1, row - 8)
        ws.cell(row, 10, f"=IF(OR(H{row}<>\"\",I{row}<>\"\"),H{row}+I{row},\"\")")
        for col in range(1, 14):
            c = ws.cell(row, col)
            c.font = Font(name="Arial", size=10, color=TEXT)
            c.alignment = Alignment(horizontal="center" if col in [1, 2, 3, 8, 9, 10, 11, 12] else "left", vertical="center", wrap_text=True)
            set_cell_fill(c, "FFFFFF")
            add_border(c)
        ws.cell(row, 8).number_format = '#,##0'
        ws.cell(row, 9).number_format = '#,##0'
        ws.cell(row, 10).number_format = '#,##0'

    summary_row = 30
    ws.merge_cells(start_row=summary_row, start_column=1, end_row=summary_row, end_column=7)
    ws.cell(summary_row, 1, "合計")
    ws.cell(summary_row, 8, "=SUM(H9:H28)")
    ws.cell(summary_row, 9, "=SUM(I9:I28)")
    ws.cell(summary_row, 10, "=SUM(J9:J28)")
    for col in range(1, 14):
        c = ws.cell(summary_row, col)
        c.font = Font(name="Arial", bold=True, color="7C3F00")
        c.alignment = Alignment(horizontal="center", vertical="center")
        set_cell_fill(c, LIGHT_ORANGE)
        add_border(c)
        if col in [8, 9, 10]:
            c.number_format = '#,##0'

    sign_row = 33
    sign_labels = ["申請人", "保管人", "直屬主管", "行政部門主任", "會計"]
    for i, label in enumerate(sign_labels):
        start = 1 + i * 2
        ws.merge_cells(start_row=sign_row, start_column=start, end_row=sign_row, end_column=start + 1)
        ws.merge_cells(start_row=sign_row + 1, start_column=start, end_row=sign_row + 2, end_column=start + 1)
        ws.cell(sign_row, start, label)
        ws.cell(sign_row, start).font = Font(name="Arial", bold=True, color="7C3F00")
        ws.cell(sign_row, start).alignment = Alignment(horizontal="center", vertical="center")
        for r in range(sign_row, sign_row + 3):
            for c in range(start, start + 2):
                add_border(ws.cell(r, c))

    type_dv = DataValidation(type="list", formula1='"收據,統一發票,免用統一發票收據,其他"', allow_blank=True)
    pay_dv = DataValidation(type="list", formula1='"現金,銀行轉帳,公司卡,其他"', allow_blank=True)
    proof_dv = DataValidation(type="list", formula1='"是,否"', allow_blank=True)
    ws.add_data_validation(type_dv)
    ws.add_data_validation(pay_dv)
    ws.add_data_validation(proof_dv)
    type_dv.add("C9:C28")
    pay_dv.add("K9:K28")
    proof_dv.add("L9:L28")

    ws.freeze_panes = "A9"
    ws.auto_filter.ref = "A8:M28"
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.print_title_rows = "8:8"
    ws.page_margins.left = 0.35
    ws.page_margins.right = 0.35
    ws.page_margins.top = 0.45
    ws.page_margins.bottom = 0.45

    wb.save(XLSX_PATH)


def shade_docx_cell(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_docx_cell_text(cell, text, bold=False, color=TEXT, size=9, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("零用金明細表")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(TEXT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("請逐筆填寫支出明細，並將憑證照片或電子檔一併上傳系統送簽核。")
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("6B7280")

    info = doc.add_table(rows=2, cols=8)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.style = "Table Grid"
    labels = ["申請日期", "", "公司 / 法人", "", "組別 / 部門", "", "保管人", ""]
    labels2 = ["申請類型", "", "上期餘額", "", "核定額度", "", "本次申請補足金額", ""]
    for row_idx, values in enumerate([labels, labels2]):
        for col_idx, value in enumerate(values):
            cell = info.cell(row_idx, col_idx)
            if col_idx % 2 == 0:
                shade_docx_cell(cell, LIGHT_ORANGE)
                set_docx_cell_text(cell, value, bold=True, color="7C3F00", size=9)
            else:
                set_docx_cell_text(cell, value, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()

    headers = ["序號", "日期", "憑證類型", "憑證號碼", "費用類別", "摘要 / 用途", "付款對象", "未稅", "營業稅", "含稅金額", "支付方式", "已附憑證", "備註"]
    table = doc.add_table(rows=12, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for col, header in enumerate(headers):
        cell = table.cell(0, col)
        shade_docx_cell(cell, ORANGE)
        set_docx_cell_text(cell, header, bold=True, color="FFFFFF", size=7)
    for row in range(1, 11):
        for col in range(len(headers)):
            cell = table.cell(row, col)
            set_docx_cell_text(cell, str(row) if col == 0 else "", size=7, align=WD_ALIGN_PARAGRAPH.CENTER if col in [0, 1, 2, 7, 8, 9, 10, 11] else WD_ALIGN_PARAGRAPH.LEFT)
    for col in range(len(headers)):
        cell = table.cell(11, col)
        shade_docx_cell(cell, LIGHT_ORANGE)
        set_docx_cell_text(cell, "合計" if col == 5 else "", bold=True, color="7C3F00", size=8)

    doc.add_paragraph()

    sign = doc.add_table(rows=2, cols=5)
    sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign.style = "Table Grid"
    for col, label in enumerate(["申請人", "保管人", "直屬主管", "行政部門主任", "會計"]):
        shade_docx_cell(sign.cell(0, col), LIGHT_ORANGE)
        set_docx_cell_text(sign.cell(0, col), label, bold=True, color="7C3F00", size=9)
        set_docx_cell_text(sign.cell(1, col), "\n\n", size=9)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    nr = note.add_run("備註：含稅金額為系統認列與補足零用金的主要金額；未稅與營業稅請依憑證內容填寫。")
    nr.font.name = "Arial"
    nr.font.size = Pt(8)
    nr.font.color.rgb = RGBColor.from_string("6B7280")

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_xlsx()
    build_docx()
    print(XLSX_PATH)
    print(DOCX_PATH)
