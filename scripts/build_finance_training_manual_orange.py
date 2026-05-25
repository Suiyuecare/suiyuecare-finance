from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "歲悅會計系統教育訓練手冊_橘色版.docx"
LOGO = ROOT / "assets" / "suiyue-logo-transparent.png"


BLUE = RGBColor(217, 119, 6)
DARK_BLUE = RGBColor(146, 64, 14)
INK = RGBColor(30, 41, 59)
MUTED = RGBColor(100, 116, 139)
HEADER_FILL = "FFEDD5"
LIGHT_FILL = "FFF7ED"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_cell_text(cell, bold=False, color=None, size=9.5):
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(size)
            run.font.bold = bold
            if color:
                run.font.color.rgb = color


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.allow_autofit = False
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, HEADER_FILL)
        style_cell_text(cell, bold=True, color=INK, size=9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            style_cell_text(cells[i], size=9.3)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_note(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.text = ""
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    r.font.size = Pt(10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    run = p2.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = INK
    set_table_width(table, [6.5])
    doc.add_paragraph()


def add_page_break_before(paragraph):
    paragraph.paragraph_format.page_break_before = True


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = DARK_BLUE
    title.paragraph_format.space_after = Pt(6)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(12)

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.color.rgb = BLUE
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.color.rgb = BLUE
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.color.rgb = DARK_BLUE
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)

    for list_name in ["List Bullet", "List Number"]:
        st = styles[list_name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25
        st.paragraph_format.left_indent = Inches(0.375)
        st.paragraph_format.first_line_indent = Inches(-0.188)

    header = section.header.paragraphs[0]
    header.text = "歲悅 Finance OS 會計系統教育訓練手冊"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = DARK_BLUE

    footer = section.footer.paragraphs[0]
    footer.text = "內部教育訓練文件｜請勿外流"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = DARK_BLUE


def build_doc():
    doc = Document()
    setup_styles(doc)

    if LOGO.exists():
        logo = doc.add_paragraph()
        logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo.add_run()
        logo_run.add_picture(str(LOGO), width=Inches(1.15))

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("歲悅 Finance OS\n會計系統教育訓練手冊")
    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("適用對象：一般員工、主管、行政部門主任、總務、人資、會計、執行長｜版本 v1.0｜建議訓練時間 60 至 90 分鐘")
    add_note(
        doc,
        "訓練目標",
        "讓使用者能獨立完成申請、暫存、送簽、簽核、附件下載、收款確認與會計入帳查詢；同時理解支出法人、支出部門、傳票與三表之間的資料關係。",
    )

    doc.add_heading("一、系統總覽", level=1)
    doc.add_paragraph(
        "Finance OS 是歲悅內部會計申請與簽核系統，用來集中處理支出申請、發票開立、收款確認、付款入帳、傳票查詢與三表資料。系統設計原則是易使用上手、資料完整、即時同步、避免重複計算。"
    )
    doc.add_paragraph(
        "使用者不需要理解完整會計分錄才能送單，但必須正確填寫申請內容、公司法人、部門、金額與附件。會計會在最後關卡確認憑據與金流，系統才會正式產生傳票並寫入帳冊。"
    )

    doc.add_heading("二、角色與權限", level=1)
    add_table(
        doc,
        ["角色", "主要任務", "常用功能"],
        [
            ["一般員工", "建立申請、上傳憑據、查看自己的申請與待補件事項", "新增申請、放款申請、簽核管理、通知中心"],
            ["課長 / 部門主管", "檢查費用合理性、部門預算與附件完整性", "簽核管理、我的簽核、歷史紀錄"],
            ["行政部門主任", "檢查行政流程、付款或發票資料、系統設定異動", "簽核管理、系統設定、通知中心"],
            ["總務", "採購完成後補上實際支出與發票憑據", "採購流程、簽核管理"],
            ["人資", "人事費用相關審核", "人事費用申請、簽核管理"],
            ["會計", "確認憑據、付款資料、傳票與三表入帳", "傳票查詢、三表總覽、收款情形、會計科目"],
            ["執行長", "最終授權、重大費用核准、系統設定檢核", "簽核管理、儀表板、系統設定"],
        ],
        [1.35, 3.05, 2.1],
    )

    doc.add_heading("三、登入與會計工作台", level=1)
    add_numbers(
        doc,
        [
            "開啟公司提供的 Finance OS 網址。",
            "使用公司帳號登入；測試階段可使用內建測試帳號。",
            "登入後直接進入會計工作台。",
            "依左側功能選單進入儀表板、申請、簽核、傳票、報表或系統設定。",
            "若畫面沒有正確顯示權限，請登出後重新登入。",
        ],
    )
    add_note(
        doc,
        "正式網址建議",
        "正式上線建議使用 finance.suiyuecare.com，並同步設定 Google OAuth 與 Supabase Auth 的正式網址與 callback URL。",
    )

    doc.add_heading("四、申請類型與使用情境", level=1)
    add_table(
        doc,
        ["類型", "使用情境", "重點附件"],
        [
            ["費用報銷", "員工已代墊費用，向公司申請報銷", "發票、收據、憑據照片"],
            ["付款申請", "公司直接付款給廠商或外部單位", "請款單、合約、發票"],
            ["預支申請", "員工先向公司預支款項，之後核銷", "預支申請書、預計用途"],
            ["零用金申請", "部門申請零用金補足或初次申請", "一般申請需憑據；初次申請免憑據"],
            ["差旅申請", "出差前申請交通、住宿、膳雜等費用", "住宿、交通、停車或其他憑據"],
            ["採購申請", "物品、設備、服務、軟體或保險採購", "採購明細、報價、後續發票"],
            ["退費申請", "退還客戶、學員、案家或合作單位款項", "退費明細、原發票或退款證明"],
            ["人事費用申請", "薪資、獎金、勞報、保費、資遣等", "薪資表、人事費用表、勞務報酬單"],
        ],
        [1.35, 3.2, 1.95],
    )

    doc.add_heading("五、一般申請操作流程", level=1)
    add_numbers(
        doc,
        [
            "點選左側「新增申請」。",
            "選擇申請類型。",
            "填寫支出法人、支出部門、申請人、表單填寫日期。",
            "填寫申請內容與匯款資訊。",
            "上傳發票、憑據、存摺封面或相關文件。",
            "檢查即時 Excel 明細表與合計金額。",
            "可先按「暫存」；確認後按「送出申請」。",
            "到「簽核管理」確認是否已出現在「我的簽核」或「歷史紀錄」。",
        ],
    )
    add_note(
        doc,
        "資料歸屬提醒",
        "支出法人會影響公司三表，支出部門會影響部門分類帳。申請人資料只是送件者身分，不等於費用歸屬。",
    )

    doc.add_heading("六、暫存表單", level=1)
    doc.add_paragraph("每張表單都可以暫存。暫存表單只會顯示給建立者本人，其他人不應看到。")
    add_numbers(
        doc,
        [
            "在表單填寫途中按「暫存」或「儲存草稿」。",
            "回到「放款申請」或「簽核管理」中的「暫存表單」。",
            "點選「繼續填寫」可回到表單。",
            "若不再使用，可點選刪除。",
        ],
    )
    add_note(
        doc,
        "附件限制",
        "瀏覽器基於安全限制，部分本機檔案可能無法在草稿中完整恢復。正式送出前，請再次確認附件是否仍在清單中。",
    )

    doc.add_heading("七、懶人費用申請", level=1)
    add_numbers(
        doc,
        [
            "進入「放款申請」後點選懶人申請。",
            "上傳一張或多張發票或憑據照片。",
            "按 OpenAI 解析憑據。",
            "系統會讀取發票號碼、日期、品項、數量、單價、未稅、營業稅與含稅金額。",
            "右側即時 Excel 明細表可直接修改。",
            "送出後，系統會將明細表轉成可下載 Excel 附件，憑據照片會彙整成 PDF 給主管檢核。",
        ],
    )
    doc.add_paragraph("如果辨識結果不準，申請人仍須手動修正。送出前，申請人對資料正確性負責。")

    doc.add_heading("八、人事費用申請", level=1)
    doc.add_paragraph("人事費用申請適用整批薪資、獎金、勞務報酬、勞健保、勞退、補充保費與資遣費。")
    add_table(
        doc,
        ["項目", "說明", "常見附件"],
        [
            ["50薪資", "薪資整批撥款", "薪資表、人事費用支出表"],
            ["獎金", "績效、開案或專案獎金", "獎金明細、核准資料"],
            ["9A/9B勞報單", "接案人員或勞務報酬", "勞務報酬單必傳"],
            ["勞保 / 健保 / 勞退", "保費或提繳費用", "保費明細或繳款資料"],
            ["二代健保補充保費", "補充保費扣繳", "補充保費明細"],
            ["資遣費", "資遣或離職給付", "離職結算資料"],
            ["其他", "其他人事相關支出", "依實際情況上傳"],
        ],
        [1.55, 2.8, 2.15],
    )
    add_bullets(
        doc,
        [
            "申請人名字、申請人部門、申請人法人會由帳號預設帶入。",
            "支出法人與支出部門是三表與部門分類帳認列依據。",
            "人事費用明細表欄位包含員工編號、姓名、類別、薪資月份、匯款金額、銀行、分行、帳號與備註。",
            "除員工編號與備註外，其餘欄位為必填。",
            "填寫即時 Excel 明細表時，按 Enter 可新增下一列並跳到同一欄。",
        ],
    )

    doc.add_heading("九、特殊表單重點", level=1)
    add_table(
        doc,
        ["表單", "操作重點"],
        [
            ["差旅申請", "填寫差旅類型、出差事由、預期成果、時間地點；住宿、交通、停車與其他費用需上傳憑據，膳雜費可依辦法帶入。"],
            ["採購申請", "申請人先填預估金額；總務完成採購後上傳發票並填寫實際支出；會計以實際支出入帳。"],
            ["退費申請", "分為沒有憑據與發票作廢；發票作廢需上傳原發票電子檔，退費金額不得大於原收款金額。"],
            ["零用金申請", "初次申請列為暫付款；一般申請依憑據補足並認列支出。"],
        ],
        [1.4, 5.1],
    )

    doc.add_heading("十、發票開立與收款", level=1)
    doc.add_paragraph(
        "發票開立也需要走簽核流程。發票流程完成後，收入會列入帳上收入；但現金流量表只有在會計確認已收款並經行政部門主任複核後才會認列現金流入。"
    )
    add_bullets(
        doc,
        [
            "收款情形分為單張發票與整批發票。",
            "會計確認收款時，應上傳入帳截圖、銀行明細或其他收款證明。",
            "行政部門主任複核通過後，現金流量表才更新。",
        ],
    )

    doc.add_heading("十一、簽核流程", level=1)
    add_table(
        doc,
        ["關卡", "角色", "主要任務"],
        [
            ["1", "申請人送件", "確認申請內容、附件與 Excel 明細。"],
            ["2", "申請人上一層級主管", "確認業務必要性與金額合理性。"],
            ["3", "申請人部門主管", "確認部門預算與成本歸屬，若與上一層主管相同則跳過。"],
            ["4", "行政部門主任", "確認行政流程、付款資料與附件完整性。"],
            ["5", "會計", "確認會計科目、稅務處理、憑據與匯款資料。"],
            ["6", "執行長", "最終授權。"],
            ["7", "申請人確認", "確認已撥款或已取得會計開立的發票。"],
            ["8", "會計確認入帳", "產生傳票並寫入分類帳與三表。"],
        ],
        [0.55, 1.9, 4.05],
    )
    doc.add_paragraph("每一關都可以填寫備註、上傳附件、核准、退件或退回上一關。若需要加簽，該關卡通過後，加簽人會成為下一關。")

    doc.add_heading("十二、附件與下載", level=1)
    add_bullets(
        doc,
        [
            "即時 Excel 明細表產生的線下 Excel 檔。",
            "憑據照片彙整 PDF。",
            "存摺封面獨立附件。",
            "申請人、主管或會計在各簽核關卡上傳的補充文件。",
        ],
    )
    add_note(doc, "檔案原則", "存摺封面不應與憑據照片合併，應獨立提供下載，避免主管檢視憑據時混淆。")

    doc.add_heading("十三、傳票與三表", level=1)
    add_bullets(
        doc,
        [
            "公司支出的三表追溯「支出法人」。",
            "部門分類帳追溯「支出部門」。",
            "執行長核准後，現金流量表會反映實際付款金流。",
            "最後一關會計確認入帳後，資產負債表與損益表才會更新。",
            "系統不得在會計最後入帳時重複更新現金流量表。",
            "發票一開立就認列收入；收款確認只影響現金流量表。",
        ],
    )

    doc.add_heading("十四、系統設定", level=1)
    doc.add_paragraph("系統設定用來維護法人資訊、部門資訊、會計科目、每間公司的手續費金額，以及各層級可使用的功能。")
    doc.add_paragraph("行政部門主任可啟動編輯，但設定異動需要執行長檢核後才可正式通過。平常未啟動編輯時，資料應保持唯讀，避免誤改。")

    doc.add_heading("十五、常見問題", level=1)
    add_table(
        doc,
        ["問題", "處理方式"],
        [
            ["主管看不到待簽核", "確認申請單是否已出現在我的簽核或歷史紀錄；若待簽核數字有增加但列表沒有內容，請重新整理並提供申請單號。"],
            ["暫存失敗", "確認目前是登入狀態；附件太大或瀏覽器阻擋檔案恢復時，請重新上傳附件後再送出。"],
            ["OpenAI 辨識失敗", "檢查 API 額度、Edge Function、OPENAI_API_KEY 與圖片清晰度；必要時手動填寫明細。"],
            ["附件下載不了", "檢查 Supabase Storage、Storage RLS 與下載 URL。主管至少應能下載 Excel 明細與憑據彙整 PDF。"],
            ["三表數字異常", "先檢查支出法人、支出部門與傳票科目；三表只應使用真實 ledger / voucher 資料。"],
        ],
        [1.85, 4.65],
    )

    doc.add_heading("十六、上線前驗收清單", level=1)
    add_table(
        doc,
        ["檢查項目", "合格標準"],
        [
            ["登入", "可使用測試帳號或公司帳號登入。"],
            ["權限", "不同角色看到的功能不同，且符合權限設定。"],
            ["暫存", "暫存後只顯示給本人。"],
            ["送出申請", "送出後出現在我的簽核與歷史紀錄。"],
            ["待我簽核", "下一關人員看得到待辦並可開啟明細。"],
            ["附件", "Excel、PDF、存摺封面可下載。"],
            ["加簽", "原關卡通過後，加簽人成為下一關。"],
            ["退回上一關", "可填備註並回到上一關。"],
            ["傳票", "會計最後確認後產生傳票。"],
            ["三表", "法人、部門、收入、支出、現金流沒有重複計算。"],
            ["OCR", "發票與存摺封面可辨識或可手動修正。"],
            ["Supabase", "資料表、Storage、Edge Function、RLS 均正常。"],
        ],
        [1.65, 4.85],
    )

    doc.add_heading("十七、建議訓練流程", level=1)
    add_table(
        doc,
        ["時間", "內容"],
        [
            ["10 分鐘", "系統目標、角色與權限介紹"],
            ["15 分鐘", "員工新增申請、暫存、附件上傳示範"],
            ["15 分鐘", "主管簽核、退回、備註、加簽示範"],
            ["15 分鐘", "會計確認、傳票、收款、三表示範"],
            ["10 分鐘", "人事費用、採購、差旅等特殊表單說明"],
            ["10 分鐘", "常見問題與正式上線注意事項"],
        ],
        [1.4, 5.1],
    )

    doc.add_heading("十八、上線後管理建議", level=1)
    add_bullets(
        doc,
        [
            "第 1 週先以試營運方式上線，每日檢查申請、簽核、傳票與三表。",
            "每週由會計檢查傳票與附件是否完整。",
            "每月結帳前，先匯出或備份當月資料。",
            "權限異動與系統設定異動需保留 audit log。",
            "API key、service role key、GitHub token 不得貼在前端或公開文件中。",
            "員工若不確定申請類型，應先詢問會計或行政部門主任再送出。",
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
